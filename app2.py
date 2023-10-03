#-*- coding: utf-8 -*-
import random,os,json,io,re,zipfile,tempfile
import ssl
import numpy as np
import pandas as pd
import streamlit as st
import streamlit_toggle as tog
from langchain import HuggingFaceHub
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs,convert_image_to_searchable_pdf
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
from huggingface_hub import login
import cv2
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from fpdf import FPDF


#from playsound import playsound
#from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.chains.summarize import load_summarize_chain
#import os
#import pyaudiogi
#import wave
#from langchain.document_loaders import UnstructuredPDFLoader
#import streamlit.components.v1 as components
#from st_custom_components import st_audiorec, text_to_docs
#import sounddevice as sd
#from scipy.io.wavfile import write

# Setting Env
if st.secrets["OPENAI_API_KEY"] is not None:
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
else:
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY")

# os.environ["OPENAI_API_KEY"] = api_key

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    return output_pdf

def new_line(string):
    pass

@st.cache_data
def usellm(prompt):
    """
    Getting GPT-3.5 Model into action
    """
    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a fraud analyst, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

# Setting Config for Llama-2
login(token=st.secrets["HUGGINGFACEHUB_API_TOKEN"])
os.environ["HUGGINGFACEHUB_API_TOKEN"] = st.secrets["HUGGINGFACEHUB_API_TOKEN"]



llama_13b = HuggingFaceHub(
            repo_id="meta-llama/Llama-2-13b-chat-hf",
            model_kwargs={"temperature":0.01, 
                        "min_new_tokens":100, 
                        "max_new_tokens":500})

memory = ConversationSummaryBufferMemory(llm= llama_13b, max_token_limit=500)
conversation = ConversationChain(llm= llama_13b, memory=memory,verbose=False)


@st.cache_data
def llama_llm(_llm,prompt):

    response = _llm.predict(prompt)
    return response

@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text

@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

@st.cache_data
def embedding_store(pdf_files):
    merged_pdf = merge_pdfs(pdf_files)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    text = ""
    for page in final_pdf.pages:
        text += page.extract_text()
    texts =  text_splitter.split_text(text)
    docs = text_to_docs(texts)
    docsearch = FAISS.from_documents(docs, hf_embeddings)
    return docs, docsearch
    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()


# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False



def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            all_text.append(text)
    return "\n".join(all_text)


# convert scanned pdf to searchable pdf
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text


# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []
if "tmp_table_gpt" not in st.session_state:
    st.session_state.tmp_table_gpt=pd.DataFrame()
if "tmp_table_llama" not in st.session_state:
    st.session_state.tmp_table_llama=pd.DataFrame()
if "tmp_summary_gpt" not in st.session_state:
    st.session_state["tmp_summary_gpt"] = ''
if "tmp_summary_llama" not in st.session_state:
    st.session_state["tmp_summary_llama"] = ''
if "sara_recommendation_gpt" not in st.session_state:
    st.session_state["sara_recommendation_gpt"] = ''
if "sara_recommendation_llama" not in st.session_state:
    st.session_state["sara_recommendation_llama"] = ''
if "case_num" not in st.session_state:
    st.session_state.case_num = ''
if "fin_opt" not in st.session_state:
    st.session_state.fin_opt = ''
if "context_1" not in st.session_state:
    st.session_state.context_1 = ''
if "llm" not in st.session_state:
    st.session_state.llm = 'Closed-Source'
if "pdf_files" not in st.session_state:
    st.session_state.pdf_files = []

# reading files from local directory from fetch evidence button
directoty_path = "data3/"
fetched_files = read_pdf_files(directoty_path)


# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 145px;
            height: 35px;
        }
    </style>
""", unsafe_allow_html=True)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)

    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)

@st.cache_data
def create_zip_file(file_paths, zip_file_name):
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, os.path.basename(file_path))

# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)


# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

#Adding llm type-> st.session_state.llm
st.session_state.llm = st.radio("",options = pd.Series(["Closed-Source","Open-Source"]), horizontal=True)

st.markdown("---")

st.title("Suspicious Activity Reporting Assistant")
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <style>
    .navbar-brand img {
      max-height: 50px; /* Adjust the height of the logo */
      width: auto; /* Allow the width to adjust based on the height */
    }
    </style>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="navbar-brand" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
                <strong>| Operations Process Automation</strong>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-bottom navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
        <!--p style='color: white;'><b>Powered by EXL</b></p--!>
        <p style='color: white;'> <strong>Powered by EXL</strong> </p>
            <!--a class="nav-link disabled" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a--!>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")
    st.markdown("---")

    # Add a drop-down for case type
    option1 = ["Select Case Type", "Fraud transaction dispute", "Money Laundering", "Insider Trading"]
    selected_option_case_type = st.sidebar.selectbox("", option1)
    st.markdown("---")
    
    # Add a single dropdown
    option2 = ["Select Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", option2)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (3).png", use_column_width=True)

    
# Assing action to the main section
if selected_option_case_type == "Select Case Type":
    st.header("")
elif selected_option_case_type == "Fraud transaction dispute":
    st.markdown("### :blue[Fraud transaction dispute]")
elif selected_option_case_type == "Money Laundering":
    st.markdown("### :red[Money Laundering]")
elif selected_option_case_type == "Insider Trading":
    st.markdown("### :red[Insider Trading]")
# st.markdown('---')

# Redirect to Merge PDFs page when "Merge PDFs" is selected
if selected_option == "SAR-2023-24680":
    st.session_state.case_num = "SAR-2023-24680"
    # st.header("Merge Documents")
    # st.write("Upload multiple document files and merge them into one doc.")

    # Upload PDF files
    # st.subheader("Upload Case Files")
    # st.markdown(f"**Case No: {st.session_state.case_num}**")
    # st.markdown("""
    #     | Case No.                  | Case Type                 | Customer Name             | Case Status             | Open Date              |
    #     | ------------------------  | ------------------------- | ------------------------- | ------------------------|------------------------|
    #     | SAR-2023-24680            | Fraud Transaction Dispute | John Brown                | In Progress             | 12/10/2020             |
    #     """)

    col1,col2 = st.columns(2)
    # Row 1
    with col1:
        st.markdown("##### **Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
        st.markdown("##### **Customer name  :** John Brown")


    with col2:
        st.markdown("##### **Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Feb 02, 2021")
        st.markdown("##### **Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Fraud transaction")


    # Row 2
    with col1:
        st.markdown("##### **Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 9659754")


    with col2:
        st.markdown("##### **Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")

st.markdown("---")


####### This markdown is to manage app style
st.markdown("""

<style>

.st-d5 {
    line-height: 1;
}


.css-1upf7v9 { 
    gap: 0.5rem;
}

.css-1balh2r{
    gap: 0;
}

.css-1544g2n {
    padding: 0;
    padding-top: 2rem;
    padding-right: 1rem;
    padding-bottom: 1.5rem;
    padding-left: 1rem;
}

.css-1q2g7hi {
    top: 2px;
    min-width: 350px;
    max-width: 600px;
    }

.st-ah {
    line-height: 1;
}

.st-af {
    font-size: 1.5rem;
}

.css-1a65djw {
    gap: 0;
    }

.css-1y4p8pa {
    width: 100%;
    padding: 3rem 1rem 10rem;
    padding-top: 3rem;
    padding-bottom: 10rem;
    max-width: 60rem;
}

.css-xujc5b p{
   font-size: 25px;
}

.css-jzprzu {
    height: 2rem;
    min-height: 1rem;
    }

</style>
""", unsafe_allow_html=True)

col1_up, col2_up, col3_up, col4_up, col5_up = st.tabs(["Data", "Generate Insights","Summarization","Download Report", "Make a Decision"])

with col1_up:
    bt1_up, bt2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])

    with bt1_up:
        # Set the color
        # st.markdown(
        #     """
        #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
        #         <span style="font-size: 16px;  ">Fetch Evidence</span>
        #     </div>
        #     """,
        #     unsafe_allow_html=True
        # )
        if 'clicked' not in st.session_state:
            st.session_state.clicked = False
        
        def set_clicked():
            st.session_state.clicked = True
            st.session_state.disabled = True
        st.write("") #for the gap
        st.button('Fetch Evidence', on_click=set_clicked)

        if st.session_state.clicked:
            # st.write("Evidence Files:") 
            # st.markdown(html_str, unsafe_allow_html=True)
            
            # Showing files
            # show_files = fetched_files.copy()
            # show_files = show_files + ['Other.pdf']
            # files_frame = pd.DataFrame(show_files, columns=["File Name"])
            # # files_frame["Select"] = [True for _ in range(len(files_frame))]
            # files_frame = files_frame.reset_index(drop=True)

            # # Add checkboxes to the DataFrame
            # df_with_checkboxes = add_checkboxes_to_dataframe(files_frame)
            
            # # Iterate through each row and add checkboxes
            # for index, row in df_with_checkboxes.iterrows():
            #     if index < len(df_with_checkboxes) - 1:
            #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
            #         df_with_checkboxes.loc[index, 'Select'] = checkbox_state
            #     else:
            #         st.checkbox(f"{row['File Name']}", value=False)



            # st.dataframe(files_frame)
            # st.write(df_reset.to_html(index=False), unsafe_allow_html=True)
            # st.markdown(files_frame.style.hide(axis="index").to_html(), unsafe_allow_html=True)
            
            
            
            #select box to select file
            selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
            st.write("Selected File: ", selected_file_name)
            st.session_state.disabled = False
            file_ext = tuple("pdf")
            if selected_file_name.endswith(file_ext):
                selected_file_path = os.path.join(directoty_path, selected_file_name)
                #converting pdf data to bytes so that render_pdf_as_images could read it
                file = pdf_to_bytes(selected_file_path)
                pdf_images = render_pdf_as_images(file)
                #showing content of the pdf
                st.subheader(f"Contents of {selected_file_name}")
                for img_bytes in pdf_images:
                    st.image(img_bytes, use_column_width=True)
            else:
                selected_file_path = os.path.join(directoty_path, selected_file_name)
                # This is showing png,jpeg files
                st.image(selected_file_path, use_column_width=True)



    with bt2_up:
        pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
        st.session_state.pdf_files = pdf_files
        # showing files
        for uploaded_file in pdf_files:
            #This code is to show pdf files
            file_ext = tuple("pdf")
            if uploaded_file.name.endswith(file_ext):
                # Show uploaded files in a dropdown
                # if pdf_files:
                st.subheader("Uploaded Files")
                file_names = [file.name for file in pdf_files]
                selected_file = st.selectbox(":blue[Select a file]", file_names)
                # Enabling the button
                st.session_state.disabled = False
                # Display selected PDF contents
                if selected_file:
                    selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                    pdf_images = render_pdf_as_images(selected_pdf)
                    st.subheader(f"Contents of {selected_file}")
                    for img_bytes in pdf_images:
                        st.image(img_bytes, use_column_width=True)

            else:
                # This is showing png,jpeg files
                st.image(uploaded_file, use_column_width=True)

#creating temp directory to have all the files at one place for accessing

    
    tmp_dir_ = tempfile.mkdtemp()
    temp_file_path= []


    for uploaded_file in pdf_files:
        file_ext = tuple("pdf")
        if uploaded_file.name.endswith(file_ext):
            file_pth = os.path.join(tmp_dir_, uploaded_file.name)
            with open(file_pth, "wb") as file_opn:
                file_opn.write(uploaded_file.getbuffer())
                temp_file_path.append(file_pth)
        else:
            pass


    for fetched_pdf in fetched_files:
        file_ext = tuple("pdf")
        if fetched_pdf.endswith(file_ext):
            file_pth = os.path.join(directoty_path, fetched_pdf)
            # st.write(file_pth)
            temp_file_path.append(file_pth) 
        else:
            pass   

     #Adding pytesseract here
    # To convert generated to pdf and save in temp direc.
    def create_pdf(text,file_name):
        # Create a new FPDF object
        pdf = FPDF()
        # Add a new page to the PDF
        pdf.add_page()
        # Set the font and font size
        pdf.set_font('Arial', size=12)
        pdf.cell(200, 10, txt=text.encode('utf-8').decode('latin-1'), ln = 1, align = 'C')
        # Write the text to the PDF
        # pdf.write(text.encode('utf-8').decode('latin-1'), ln = 1, align = 'C')
        # Save the PDF
        pdf.output(os.path.join(tmp_dir_,file_name))
        file_pth = os.path.join(tmp_dir_,file_name)
        temp_file_path.append(file_pth)

    # # # Pytesseract code 

    # #file path for uploaded files
    # file_pth = []
    # for uploaded_file in pdf_files:
    #     file_ext1 = tuple("pdf")
    #     file_ext2 = tuple(["png","jpeg"])
    #     if uploaded_file.name.endswith(file_ext1):
    #         file_pth_= os.path.join(tmp_dir_, uploaded_file.name)
    #         with open(file_pth_, "wb") as file_opn:
    #             file_opn.write(uploaded_file.getbuffer())
    #             file_pth.append(file_opn)
    #     elif uploaded_file.name.endswith(file_ext2):
    #         file_pth_= os.path.join(tmp_dir_, uploaded_file.name)
    #         file_pth.append(file_pth_)
    #     else:
    #         pass

 

    # # For uploaded files
    # for file in file_pth:
    #     st.write(file.name)
    #     file_ext1 = tuple("pdf")
    #     file_ext2 = tuple(["png","jpeg"])
    #     if file.endswith(file_ext1):
    #         if is_searchable_pdf(file)==False:
    #             text = convert_scanned_pdf_to_searchable_pdf(file)
    #             create_pdf(text,'uploaded_file.pdf')
    #         else:
    #             with open(file, "wb") as file_opn:
    #                 file_opn.write(file.getbuffer())
    #                 temp_file_path.append(file_opn)           
    #     elif file.endswith(file_ext2):
    #         text = convert_image_to_searchable_pdf(file)
    #         create_pdf(text,'uploaded_file.pdf') 
    #     else:
    #         pass          
       
       
    # #for fetched files
    # for fetched_pdf in fetched_files:
    #     file_ext1 = tuple("pdf")
    #     file_ext2 = tuple(["png","jpeg"])
    #     if fetched_pdf.endswith(file_ext1):
    #         selected_file_path = os.path.join(directoty_path, fetched_pdf)
    #         if is_searchable_pdf(selected_file_path)==False:
    #             text = convert_scanned_pdf_to_searchable_pdf(selected_file_path)
    #             file_name = os.path.basename(selected_file_path)
    #             split_name = file_name.split('.')
    #             create_pdf(text,f'{split_name[0]}.pdf')
    #         else:
    #             file_pth = os.path.join(directoty_path, fetched_pdf)
    #             temp_file_path.append(file_pth)
    #     elif fetched_pdf.endswith(file_ext2):
    #         selected_file_path = os.path.join(directoty_path, fetched_pdf)
    #         text = convert_image_to_searchable_pdf(selected_file_path)
    #         file_name = os.path.basename(selected_file_path)
    #         split_name = file_name.split('.')
    #         create_pdf(text,f'{split_name[0]}.pdf')

    #     else:
    #         pass
      
 
    # st.write(temp_file_path)


with col2_up:
         #This is the embedding model
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    # model_name = "hkunlp/instructor-large"
    
    # Memory setup for gpt-3.5
    llm = ChatOpenAI(temperature=0.1)
    memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
    conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
    
    
    # Adding condition on embedding
    try:
        if temp_file_path:
            hf_embeddings = embed(model_name) 
        else:
            pass
    except NameError:
        pass
    
    # Chunking with overlap
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size = 1000,
        chunk_overlap  = 100,
        length_function = len,
        separators=["\n\n", "\n", " ", ""]
    )
  
    
    try:
        if temp_file_path:
            docs, docsearch = embedding_store(temp_file_path)
        else:
            pass
    except Exception:
        pass

       

# Creating header
    col1,col2 = st.columns(2)
    with col1:
        st.markdown("""<span style="font-size: 24px; ">Key Questions</span>""", unsafe_allow_html=True)
        # st.subheader('Key Questions')
        # Create a Pandas DataFrame with your data
        data = {'Questions': [" What is the victim's name?","What is the suspect's name?",' List the merchant name',' How was the bank notified?',' When was the bank notified?',' What is the fraud type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?',' What type of cards are involved?',' Was the police report filed?']}
        df_fixed = pd.DataFrame(data)
        df_fixed.index = df_fixed.index +1
    with col2:
        # Create a checkbox to show/hide the table
        cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
        with cols1:
            show_table = tog.st_toggle_switch(label="", 
                                key="Key1", 
                                default_value=False, 
                                label_after = False, 
                                inactive_color = '#D3D3D3', 
                                active_color="#11567f", 
                                track_color="#29B5E8"
                                )
        # Show the table if the checkbox is ticked
        if show_table:
            # st.write(df_fixed)
            # st.dataframe(df_fixed, width=1000)
            df_fixed["S.No."] = df_fixed.index
            df_fixed = df_fixed.loc[:,['S.No.','Questions']]
            st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)

    with st.spinner('Wait for it...'):
          if st.button("Generate Insights",disabled=st.session_state.disabled):
                if st.session_state.llm == "Closed-Source":
                    queries ="Please provide the following information regarding the possible fraud case: What is the name of the customer name,\
                    has any suspect been reported, list the merchant name, how was the bank notified, when was the bank notified, what is the fraud type,\
                    when did the fraud occur, was the disputed amount greater than 5000 USD, what type of cards are involved, was the police report filed,\
                    and based on the evidence, is this a suspicious activity(Summarize all the questions asked prior to this in a detailed manner),that's the answer of\
                    whether this is a suspicious activity\
                    "
                    contexts = docsearch.similarity_search(queries, k=5) 
                    prompts = f" Give a the answer to the below questions as truthfully and in as detailed in the form of sentences\
                    as possible as per given context only,\n\n\
                            What is the victim's name?\n\
                            What is the suspect's name?\n\
                            List the merchant name\n\
                            How was the bank notified?\n\
                            When was the bank notified?\n\
                            What is the fraud type?\n\
                            When did the fraud occur?\n\
                            Was the disputed amount greater than 5000 USD?\n\
                            What type of cards are involved?\n\
                            Was the police report filed?\n\
                        Context: {contexts}\n\
                        Response (in the python dictionary format\
                        where the dictionary key would carry the questions and its value would have a descriptive answer to the questions asked): "
                        
                    response = usellm(prompts)

                    try:
                        resp_dict_obj = json.loads(response)
                        res_df_gpt = pd.DataFrame(resp_dict_obj.items(), columns=['Question','Answer'])
                    except:
                        e = Exception("")
                        st.exception(e)

                               

                    # try:
                        # res_df_gpt.Question = res_df_gpt.Question.apply(lambda x: x.split(".")[1])
                        # res_df_gpt.index = res_df.index + 1
                        # df_base_gpt = res_df_gpt.copy(deep=True)
                        # df_base_gpt["S.No."] = df_base_gpt.index
                        # df_base_gpt = df_base_gpt.loc[:,['S.No.','Question','Answer']]
                        # st.write(df_base_gpt)
                    # except IndexError:
                    #     pass
                    # #st.table(res_df_gpt)
                    # st.markdown(df_base_gpt.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                    # st.session_state["tmp_table_gpt"] = pd.concat([st.session_state.tmp_table_gpt, res_df_gpt], ignore_index=True)
                    
                    try:
                        res_df_gpt.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                        res_df_gpt = res_df_gpt.set_index([index_])   
                    except IndexError: 
                        pass

                    st.table(res_df_gpt)
                    st.session_state["tmp_table_gpt"] = pd.concat([st.session_state.tmp_table_gpt, res_df_gpt], ignore_index=True)


                    #SARA Recommendation
                    # query = "Is this a Suspicious Activity?"
                    # context_1 = docsearch.similarity_search(query, k=5)
                    # prompt = f'''Act as a financial analyst and give concise answer to the question as truthfully as possible with given Context.
                    #             If transaction/disputed amount is below the $5000 value threshold than there is no suspicious activity in this case. But,
                    #             If transaction/disputed amount is above the $5,000 value threshold than check for below point to make sure if it is a suspicious activity or not-
                    #             1. There is an indication of suspicion with involvement of multiple individuals whose details mismatch with customer details. (Customer details can be identified from Cardholder Information)
                    #             2. A potential suspect is identified.
                    #             Analyse above points properly and answer if there is any fraud/suspicious activity happening.\n\n\
                    #             Question: {query}\n\
                    #             Context: {context_1}\n\                      
                    #             Response: (Give me your response in pointers.)''' 
                    

                    queries ="Please provide the following information from the context: If transaction,disputed amount is above the $5000 threshold,\
                              There is an indication of suspicion with involvement of multiple individuals whose details mismatch with customer details. (Customer details can be identified from Cardholder Information),\
                              A potential suspect is identified, Mention of an individual/suspect whose details such as name and address mismatch with customer details and based on the evidence, is this a suspicious activity (Summarize all the questions asked prior to this in a detailed manner),\
                              that is the answer of whether this is a suspicious activity"
                    
                    contexts = docsearch.similarity_search(queries, k=5) 
                    prompt = f" You are professional Fraud Analyst. Find answer to the questions as truthfully and in as detailed as possible as per given context only,\n\n\
                        1. If the transaction/disputed amount > 5,000 USD value threshold.\n\n\
                        2. Is there mention of an individual/suspect whose details such as (name,address) mismatch with customer details. (Customer details can be identified from cardholder's information).\n\n\
                        3. If a potential suspect is identified.\n\n\
                        Based the above findings, identify if this can be consider as Suspicious Activity or not.\n\n\
                        If transaction/disputed amount is < 5000 USD threshold and no suspicious activity is detected based on above mentioned points, write your response as - There is no indication of suspicious activity.\n\n\
                        Context: {contexts}\n\
                        Response (Give response in pointers based on your analysis.) "
                         
                                         
                    response1 = usellm(prompt) 


                    st.session_state["sara_recommendation_gpt"] = response1                
                    
                    st.markdown("### SARA Recommendation")
                    st.markdown(response1)


                    
                elif st.session_state.llm == "Open-Source":

                    chat_history = {}

                    query = "What is the victim's name?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''You are a professional fraud analyst. Perform Name Enitity Recognition to identify the victim's name as accurately as possible, given the context. The victim can also be referenced as the customer with whom the Fraud has taken place.
                    victim's name is the Name provided in Cardholder Information.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\
                            Response: (Give me response in one sentence. Do not give me any Explanation or Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "What is the suspect's name?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''Act as a professional fraud analyst.You need to check the document and compare if any name discrepencies are present that points towards the suspect who used the card without the consent of the cardholder.
                                Take the provided information as accurate. Reply the name of the person who is the suspect. \n\n\
                                Context: {context_1}\n\
                                Response: (Give a short response in a single sentence.Do not add any extra Information, Explanation,Note.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
                    
                    
                    query = "list the merchant name"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''You are a professional fraud analyst, perform Name Enitity Recognition to identify Merchant as accurately as possible from the provided information.A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and short response.\n\n\
                      Take the provided information as accurate.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give a short response in a single sentence. Do not add any extra Information,Explanation,Note.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "How was the bank notified?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence. Do not give me any further Explanation, Note )'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response

                    
                    query = "When was the bank notified?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to identify when the bank was notified of the Fraud. Look for the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not add any prefix like 'Response' or 'Based on the document'. Do not add any extra Explanation, Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response               


                    query = "What is the Fraud Type?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me response in one sentence. Do not add prefix like 'Response' or 'based on the document. Do not give me any Explanation or Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "When did the fraud occur?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence. Do not add prefix like 'based on the document. Do not add any further Explanation or Note.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "Was the disputed amount greater than 5000 usd?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the disputed amount.Perform a mathematical calculation to identify if the disputed amount is greater than 5000 USD or not.Given the context, give a relevant and concise response.\n\n\
                                     Take the provided information as accurate. \n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give a short response in a single sentence. Do not give any extra Explanation, Note, Descricption, Information.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "What type of cards are involved?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not add prefix like: ['based on the document']. Do not add any further Explanation, Note.')'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    query = "was the police report filed?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response in a single sentence. Do not write any extra [Explanation, Note, Descricption].)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response


                    try:
                        res_df_llama = pd.DataFrame(list(chat_history.items()), columns=['Question','Answer'])
                        res_df_llama.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                        res_df_llama = res_df_llama.set_index([index_])
                        # st.write(res_df_llama)
                    except IndexError: 
                        pass
                    st.table(res_df_llama)
                    st.session_state["tmp_table_llama"] = pd.concat([st.session_state.tmp_table_llama, res_df_llama], ignore_index=True)
                
                    ## SARA Recommendation
                    query = "Is this a Suspicious Activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''Act as a financial analyst and give concise answer to the question, with given Context.
                    This can be addressed as a suspicious activity based on [transaction amount,fraud type,suspect name not matching with the customer name, suspect address does not match with the customer address].\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\                      
                                Response: (Give me a concise response in pointers)'''
                    response1 = llama_llm(llama_13b,prompt_1) 
                    st.session_state["sara_recommendation_llama"] = response1                    

                    st.markdown("### SARA Recommendation")
                    st.markdown(response1)


    st.markdown("---")

    # Text Input
    # st.markdown("""<span style="font-size: 24px; ">Ask Additional Questions</span>""", unsafe_allow_html=True)
    query = st.text_input(':black[Ask Additional Questions]',disabled=st.session_state.disabled)
    text_dict = {}
    @st.cache_data
    def LLM_Response():
        llm_chain = LLMChain(prompt=prompt, llm=llm)
        response = llm_chain.run({"query":query, "context":context})
        return response
    if st.session_state.llm == "Closed-Source":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)

                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                if query.lower() == "what is the victim's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "what is the suspect's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "list the merchant name":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "how was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "when was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "what type of fraud is taking place?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                
                elif query.lower() == "when did the fraud occur?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                        
                elif query.lower() == "was the disputed amount greater than 5000 usd?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "what type of cards are involved?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                elif query.lower() == "was the police report filed?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                        
                elif query.lower() == "Is this a valid SAR case?":
                    prompt_1 = f''' You need to act as a Financial analyst to check if this is a SAR or not, given the following context, if the transaction amount is less than 5000 USD we cannot categorize this as SAR (Suspicious activity Report).Give a relevant and concise response. \n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''

                    
                else:
                    prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\                      
                                Response: '''


                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                response = usellm(prompt_1) #LLM_Response()
                text_dict[query] = response
                # resp_dict_obj.update(text_dict)
                st.write(response)
                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()

                st.session_state["tmp_table_gpt"] = pd.concat([st.session_state.tmp_table_gpt, df], ignore_index=True)
                st.session_state.tmp_table_gpt.drop_duplicates(subset=['Question'])


    elif st.session_state.llm == "Open-Source":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)

                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                if query.lower() == "what is the victim's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.
                                Customer/Victim is cardholder, whose card is used without their consent.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response.) '''

                    
                elif query.lower() == "what is the suspect's name?":
                    prompt_1 = f''''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me response in one sentence.Do not give me any Explanation or Note)'''


                    
                elif query.lower() == "list the merchant name":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                    
                elif query.lower() == "how was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Provide a concise Response.) '''

                    
                elif query.lower() == "when was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response.)'''

                    
                elif query.lower() == "what type of fraud is taking place?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                
                elif query.lower() == "when did the fraud occur?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response..
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                        
                elif query.lower() == "was the disputed amount greater than 5000 usd?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.
                                Kindly do not provide any extra [Explanation, Note, Description] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.) '''

                    
                elif query.lower() == "what type of cards are involved?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of Card and Card Network involved, given the context. On a higher level the card can be a Dedit, Crebit Card. VISA, MasterCard, American Express, Citi Group, etc. are the different brands with respect to a Credit Card or Debit Card . Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the context.) '''

                    
                elif query.lower() == "was the police report filed?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                elif query.lower() == "is this a valid sar case?":
                    prompt_1 =  f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                Then we can address this as a valid SAR case.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise response in single sentence.Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.)'''        
                
                
                elif query.lower() == "is there any evidence of a sar case?":
                    prompt_1 = f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                Then we can address this as a SAR case.Give a concise response with the suspect name. \n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.) '''

                    
                else:
                    prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.
                                Do not provide any extra [Explanation, Note,Description] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\                      
                                Response: (Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the Response.)'''


                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                # response = usellm(prompt_1) #LLM_Response()
                response = llama_llm(llama_13b,prompt_1)
                text_dict[query] = response

                st.write(response)

                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()

                st.session_state["tmp_table_llama"] = pd.concat([st.session_state.tmp_table_llama, df], ignore_index=True)
                st.session_state.tmp_table_llama.drop_duplicates(subset=['Question'])

with col3_up:
    with st.spinner('Summarization ...'):
        st.markdown("""<span style="font-size: 24px; ">Summarize key findings of the case.</span>""", unsafe_allow_html=True)
        st.write()
        if st.button("Summarize",disabled=st.session_state.disabled):
            if st.session_state.llm == "Closed-Source":
                st.session_state.disabled=False
                summ_dict_gpt = st.session_state.tmp_table_gpt.set_index('Question')['Answer'].to_dict()
                # chat_history = resp_dict_obj['Summary']
                memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
                memory.save_context({"input": "This is the entire summary"}, 
                                {"output": f"{summ_dict_gpt}"})
                conversation = ConversationChain(
                llm=llm, 
                memory = memory,
                verbose=True)
                st.session_state["tmp_summary_gpt"] = conversation.predict(input="Provide a detailed summary of the text provided by reframing the sentences. Provide the summary in a single paragraph. Please don't include words like these: 'chat summary', 'includes information' in my final summary.")
                # showing the text in a textbox
                # usr_review = st.text_area("", value=st.session_state["tmp_summary_gpt"])
                # if st.button("Update Summary"):
                #     st.session_state["fin_opt"] = usr_review
                st.write(st.session_state["tmp_summary_gpt"])


            elif st.session_state.llm == "Open-Source":
                st.session_state.disabled=False
                template = """Write a detailed summary of the text provided.
                ```{text}```
                Response: (Return your response in a single paragraph.) """
                prompt = PromptTemplate(template=template,input_variables=["text"])
                llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)

                summ_dict_llama = st.session_state.tmp_table_llama.set_index('Question')['Answer']
                text = []
                for key,value in summ_dict_llama.items():
                    text.append(value)
                st.session_state["tmp_summary_llama"] = llm_chain_llama.run(text)
                st.write(st.session_state["tmp_summary_llama"])

    
    tmp_summary = []
    tmp_table = pd.DataFrame()   
        
    try:

        if st.session_state.llm == "Closed-Source":
            st.session_state.disabled=False
            tmp_summary.append(st.session_state["tmp_summary_gpt"])
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_gpt"]], ignore_index=True)
            tmp_table.drop_duplicates(inplace=True)
        

        elif st.session_state.llm == "Open-Source":
            st.session_state.disabled=False
            tmp_summary.append(st.session_state["tmp_summary_llama"])
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_llama"]], ignore_index=True)
            tmp_table.drop_duplicates(inplace=True)

    except: 

        e = Exception("")
        st.exception(e)

    
    try:
        # initiate the doc file
        doc = docx.Document()
        # doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(f"Case No.: {st.session_state.case_num}",0)

        # Add a subheader for case details
        subheader_case = doc.add_paragraph("Case Details")
        subheader_case.style = "Heading 2"
        # Addition of case details
        paragraph = doc.add_paragraph(" ")
        case_info = {
            "Case Number                            ": " SAR-2023-24680",
            "Customer Name                       ": " John Brown",
            "Customer ID                              ": " 9659754",
            "Case open date                         ": " Feb 02, 2021",
            "Case Type                                  ": " Fraud Transaction",
            "Case Status                                ": " Open"
        }
        for key_c, value_c in case_info.items():
            doc.add_paragraph(f"{key_c}: {value_c}")
        paragraph = doc.add_paragraph(" ")

        # Add a subheader for customer info to the document ->>
        subheader_paragraph = doc.add_paragraph("Customer Information")
        subheader_paragraph.style = "Heading 2"
        paragraph = doc.add_paragraph(" ")

        # Add the customer information
        customer_info = {
            "Name                                           ": " John Brown",
            "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
            "Phone                                          ": " (619) 425-2972",
            "A/C No.                                        ": " 4587236908230087",
            "SSN                                               ": " 653-30-9562"
        }

        for key, value in customer_info.items():
            doc.add_paragraph(f"{key}: {value}")
        paragraph = doc.add_paragraph()
        # Add a subheader for Suspect infor to the document ->>
        subheader_paragraph = doc.add_paragraph("Suspect's Info")
        subheader_paragraph.style = "Heading 2"
        paragraph = doc.add_paragraph()
        #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""

        # Add the customer information
        sent_val = "Suspect has been reported."
        paragraph = doc.add_paragraph()
        runner = paragraph.add_run(sent_val)
        runner.bold = True
        runner.italic = True
        suspect_info = {
            "Name                                             ": "Mike White",
            "Address                                        ": "520, WintergreenCt,Vancaville,CA,95587",
            "Phone                                             ": "NA",
            "SSN                                                 ": "NA",
            "Relationship with Customer  ": "NA"
        }

        for key, value in suspect_info.items():
            doc.add_paragraph(f"{key}: {value}")

        doc.add_heading('Summary', level=2)
        paragraph = doc.add_paragraph()
        doc.add_paragraph(tmp_summary)
        paragraph = doc.add_paragraph()
        doc.add_heading('Key Insights', level=2)
        paragraph = doc.add_paragraph()
        columns = list(tmp_table.columns)
        table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
        table.autofit = True
        for col in range(len(columns)):
            # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
            table.cell(0, col).text = columns[col]
        # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')
        
        for i, row in enumerate(tmp_table.itertuples()):
            table_row = table.add_row().cells # add new row to table
            for col in range(len(columns)): # iterate over each column in row and add text
                table_row[col].text = str(row[col+1]) # avoid index by adding col+1
        # save document
        # output_bytes = docx.Document.save(doc, 'output.docx')
        # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if st.session_state.llm == "Closed-Source":
            paragraph = doc.add_paragraph()
            paragraph = doc.add_paragraph()
            doc.add_heading('SARA Recommendation', level=2)
            doc.add_paragraph()       
            paragraph = doc.add_paragraph(st.session_state["sara_recommendation_gpt"])
        elif st.session_state.llm == "Open-Source":
            paragraph = doc.add_paragraph()
            paragraph = doc.add_paragraph()
            doc.add_heading('SARA Recommendation', level=2)
            doc.add_paragraph()       
            paragraph = doc.add_paragraph(st.session_state["sara_recommendation_llama"])           

        bio = io.BytesIO()
        doc.save(bio)
    except:
        e = Exception("")
        st.exception(e)
        
        

with col4_up:

    col_d1, col_d2 = st.tabs(["Download Report", "Download Case Package"])

    with col_d1:
    # Applying to download button -> download_button
        st.markdown("""
            <style>
                .stButton download_button {
                    width: 100%;
                    height: 70%;
                }
            </style>
        """, unsafe_allow_html=True)


        # combined_doc_path = os.path.join(tmp_dir, "resulting_document.docx")
        # doc.save(combined_doc_path)

        # # Create a zip file with the uploaded PDF files and the combined document
        # zip_file_name = "package_files.zip"
        # if pdf_files:
        #     st.write(file_paths)
        #     files =  [combined_doc_path]
        #     st.write(files)
            
        #     create_zip_file(files, zip_file_name)
        #     # create_zip_file(file_paths, zip_file_name)
        # else:
        #     pass
        # # Download the package
        # with open(zip_file_name, "rb") as file:
        #     st.download_button(
        #         label="Download Case Package", 
        #         data=file, 
        #         file_name=zip_file_name,
        #         disabled=st.session_state.disabled)
        
        if doc:
            st.download_button(
                label="Download Report",
                data=bio.getvalue(),
                file_name="Report.docx",
                mime="docx",
                disabled=st.session_state.disabled
            )
with col_d2:
    
    # initiating a temp file
    tmp_dir = tempfile.mkdtemp()

    file_paths= []

    for uploaded_file in st.session_state.pdf_files:
        file_pth = os.path.join(tmp_dir, uploaded_file.name)
        with open(file_pth, "wb") as file_opn:
            file_opn.write(uploaded_file.getbuffer())
        file_paths.append(file_pth)

    for fetched_pdf in fetched_files:
        # st.write(fetched_pdf)
        file_pth = os.path.join(directoty_path, fetched_pdf)
        # st.write(file_pth)
        file_paths.append(file_pth)

    
    combined_doc_path = os.path.join(tmp_dir, "report.docx")
    doc.save(combined_doc_path)



    # Create a zip file with the uploaded PDF files and the combined document
    zip_file_name = "package_files.zip"
    if file_paths:
        files =  [combined_doc_path] + file_paths
        create_zip_file(files, zip_file_name)
        # create_zip_file(file_paths, zip_file_name)
    else:
        pass

    
    # Download the package

    with open(zip_file_name, "rb") as file:
        st.download_button(
            label="Download Case Package", 
            data=file, 
            file_name=zip_file_name,
            disabled=st.session_state.disabled)

    # # Cleanup: Delete the temporary directory and its contents
    # for file_path in file_paths + [combined_doc_path]:
    #     os.remove(file_path)
    # os.rmdir(tmp_dir)

    with col5_up:   
        # Adding Radio button
             
        st.markdown("""<span style="font-size: 24px;color:#0000FF">Is SAR filing required?</span>""", unsafe_allow_html=True)

        if st.session_state.llm == "Closed-Source":
            st.write("#### *SARA Recommendation*")
            # st.markdown("""<span style="font-size: 18px;">*Based on the following findings for the underlying case, under Bank Secrecy Act, it is recommended to file this case as a suspicious activity:*</span>""", unsafe_allow_html=True)
            # st.markdown("""<span style="font-size: 18px;">*1. Transaction amount is above the $5,000 value threshold*</span>""", unsafe_allow_html=True)
            # st.markdown("""<span style="font-size: 18px;">*2. There is an indication of suspicion with involvement of multiple individuals, mismatch of customer details on merchant invoice and identification of a potential suspect*.</span>""", unsafe_allow_html=True)           
      
            query = "Give your recommendation if SAR filling is required or not?"
            context_1 = docsearch.similarity_search(query, k=5)
            prompt = f'''Act as a financial analyst and give concise answer to the question, with given Context.\n\n\
            SAR refers to Suspicious activity Report, which is a document that financial institutions must file with the Financial Crimes Enforcement Network (FinCEN) based on the Bank Secrecy Act whenever there is a suspicious activity.\n\n\
            To confirm this as a suspicious activity-
            1. The transaction/disputed amount is > 5000 USD threshold. 
            2. There is an indication of suspicion with involvement of multiple individuals whose details mismatch with customer details. (Customer details can be identified from Cardholder Information)
            3. Any potential suspect is identified.\n\n\     
            If transaction/disputed amount is < 5000 USD threshold and no suspicious activity is detected based on above mentioned points, write your response as - There is no indication of suspicious activity.Therefore,no requirement to file SAR with FinCEN.\n\n\
                    Question: {query}\n\
                    Context: {context_1}\n\                      
                    Response: (Based on your analysis give a concise response in pointers.Mention whom to file based on Bank Secrecy Act.)'''
            
            
            response1 = usellm(prompt) 
            st.markdown(f'''<em>{response1}</em>''',unsafe_allow_html=True)

        # if st.session_state.llm == "Open-Source":

        #     st.write("#### *SARA Recommendation*")

        #     query = "Is SAR filling required?"
        #     context_1 = docsearch.similarity_search(query, k=5)
        #     prompt_1 = f'''Act as a financial analyst and give concise answer to the question, with given Context.\n\n\
        #     First Find out if there is any suspicious activity based on the following-
        #     1. Transaction amount is above the $5,000 value threshold.
        #     2. There is mismatch of customer details such as name,address on merchant invoice.
        #     3. Any potential suspect is identified.\n\n\     
        #     Based on the suspicious activity, answer if SAR filling is required or not.SAR refers to Suspicious activity Report, which is a document that financial institutions must file with the Financial Crimes Enforcement Network (FinCEN) based on the Bank Secrecy Act whenever there is a suspicious activity.\n\n\       
        #             Question: {query}\n\
        #             Context: {context_1}\n\                      
        #             Response: (Give me a concise response in pointers. Also, mention whom to file based on Bank Secrecy Act.)'''
     
        #     response1 = llama_llm(llama_13b,prompt_1)
        #     st.markdown(f'''<em>{response1}</em>''',unsafe_allow_html=True)


        selected_rad = st.radio(":blue", ["Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
        if selected_rad == "Refer for review":
            email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
            email_id = st.text_input("Enter email ID")
            if email_id and not re.match(email_regex, email_id):
                st.error("Please enter a valid email ID")


        if st.button("Submit"):
            if selected_rad in ("Yes"):
                st.warning("Thanks for your review, your response has been submitted")
            elif selected_rad in ("No"):
                st.success("Thanks for your review, your response has been submitted")

            else:
                st.info("Thanks for your review, Case has been assigned to the next reviewer")


        # # Allow the user to clear all stored conversation sessions
        if st.button("Reset Session"):
            reset_session_state()
            st.cache_data.clear()
            # pdf_files.clear()
            # os.rmdir(tmp_dir_)




# Footer
st.markdown(
    """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
    """
    , unsafe_allow_html=True)
st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)


padding = 0
st.markdown(f""" <style>
    .reportview-container .main .block-container{{
        padding-top: {padding}rem;
        padding-right: {padding}rem;
        padding-left: {padding}rem;
        padding-bottom: {padding}rem;
    }} </style> """, unsafe_allow_html=True)



